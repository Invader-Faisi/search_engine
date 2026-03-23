import os
from werkzeug.utils import secure_filename
from flask import current_app
from app.models.document_model import Document
from app import db
from datetime import datetime
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

ALLOWED_EXTENSIONS = {'txt', 'csv', 'pdf', 'docx'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_file(file, user_id):
    if not file or not allowed_file(file.filename):
        return None

    filename = secure_filename(file.filename)
    project_root = os.path.dirname(current_app.root_path)
    upload_folder = os.path.join(project_root, "data", "uploads")
    os.makedirs(upload_folder, exist_ok=True)

    filepath = os.path.join(upload_folder, filename)
    file.save(filepath)

    rel_path = os.path.join("data", "uploads", filename)

    document = Document(
        filename=filename,
        filepath=rel_path,
        user_id=user_id
    )
    db.session.add(document)
    db.session.commit()

    print(f"[INFO] File saved to {filepath}, stored as {rel_path}")
    return document


def get_user_files(user_id):
    files = Document.query.filter_by(user_id=user_id).all()
    file_list = []
    for f in files:
        if os.path.exists(f.filepath):
            file_list.append({
                "name": f.filename,
                "path": f.filepath,
                "size": os.path.getsize(f.filepath),
                "date": f.upload_date.strftime("%Y-%m-%d %H:%M")
            })
    return file_list


def extract_text(filepath):
    ext = filepath.split('.')[-1].lower()

    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == "csv":
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()

        elif ext == "docx":
            doc = DocxDocument(filepath)
            return "\n".join([para.text for para in doc.paragraphs])

        elif ext == "pdf":
            reader = PdfReader(filepath)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text

    except Exception as e:
        print(f"[ERROR] Text extraction failed: {e}")
        return ""

    return ""
